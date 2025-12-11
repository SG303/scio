import asyncio
import aiofiles


async def parse_document(file_path: str, file_ext: str) -> str:
    """Parse document and extract text content"""
    
    if file_ext in [".txt", ".md"]:
        return await parse_text_file(file_path)
    elif file_ext == ".pdf":
        return await parse_pdf(file_path)
    elif file_ext in [".docx", ".doc"]:
        return await parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_ext}")


async def parse_text_file(file_path: str) -> str:
    """Parse plain text or markdown file"""
    async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:
        return await f.read()


def _sync_parse_pdf(file_path: str) -> str:
    """Synchronous PDF parsing (runs in executor to avoid blocking)"""
    from PyPDF2 import PdfReader
    
    reader = PdfReader(file_path)
    text_parts = []
    
    for page in reader.pages:
        text = page.extract_text()
        if text:
            text_parts.append(text)
    
    return "\n\n".join(text_parts)


async def parse_pdf(file_path: str) -> str:
    """Parse PDF file using run_in_executor to avoid blocking the event loop"""
    try:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, _sync_parse_pdf, file_path)
    except ImportError:
        raise ValueError("PyPDF2 is required for PDF parsing")
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")


def _sync_parse_docx(file_path: str) -> str:
    """Synchronous DOCX parsing (runs in executor to avoid blocking)"""
    from docx import Document
    
    doc = Document(file_path)
    text_parts = []
    
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text)
    
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                text_parts.append(" | ".join(row_text))
    
    return "\n\n".join(text_parts)


async def parse_docx(file_path: str) -> str:
    """Parse DOCX file using run_in_executor to avoid blocking the event loop"""
    try:
        loop = asyncio.get_running_loop()
        return await loop.run_in_executor(None, _sync_parse_docx, file_path)
    except ImportError:
        raise ValueError("python-docx is required for DOCX parsing")
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
